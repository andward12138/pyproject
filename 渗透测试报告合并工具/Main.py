from docx import Document
from docx.shared import RGBColor
import logging
import win32com.client as win32
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
import tempfile
import ctypes
import sys

from number_modifier import extract_and_modify_numbered_paragraphs, get_max_main_chapter_number_from_toc
from text_modifier import modify_first_occurrence, update_section_references
from toc_updater import update_table_of_contents
from doc_modifier import remove_end_text

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
def add_red_note_to_doc(doc_path):
    doc = Document(doc_path)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(
        "注意！交付日期，文件名，编号，插入副文档第一行标题大小以及修订记录需要手动修改，修改完成后可以删除此行提醒")
    run.font.color.rgb = RGBColor(255, 0, 0)  # 红色字体
    doc.save(doc_path)
    logging.info("红色字体的备注提醒已添加到文档末尾")

def copy_content_after_toc(doc_path):
    try:
        app = win32.Dispatch('Kwps.Application')
        app.Visible = False
        doc = app.Documents.Open(doc_path)

        temp_doc_path = os.path.join(tempfile.gettempdir(), 'temp_sub_doc.docx')
        temp_doc = app.Documents.Add()

        toc_end_position = None
        for table in doc.TablesOfContents:
            toc_end_position = table.Range.End

        if toc_end_position is not None:
            range_to_copy = doc.Range(Start=toc_end_position, End=doc.Content.End)
            range_to_copy.Copy()
            temp_doc.Range().Paste()

            temp_doc.SaveAs(temp_doc_path)
            temp_doc.Close(False)
            doc.Close(False)
            app.Quit()
            return temp_doc_path

        doc.Close(False)
        app.Quit()
        return doc_path

    except Exception as e:
        logging.error(f"复制内容时出错: {e}")
        raise

def merge_documents(main_doc_path, sub_doc_path, temp_main_doc_path, progress_callback, add_red_note=False,
                    is_last=False, temp_files=[], appendix_index=None):
    sub_doc = Document(sub_doc_path)
    if sub_doc.tables:
        logging.info("检测到目录页，复制目录之后的内容...")
        sub_doc_path = copy_content_after_toc(sub_doc_path)

    remove_end_text(main_doc_path, temp_main_doc_path, "【全文结束】")
    progress_callback(10)
    temp_files.append(temp_main_doc_path)

    update_table_of_contents(temp_main_doc_path)
    progress_callback(30)

    temp_sub_doc_path = os.path.join(os.path.dirname(temp_main_doc_path), "temp_sub_doc.docx")
    extract_and_modify_numbered_paragraphs(sub_doc_path, temp_main_doc_path, temp_sub_doc_path)
    progress_callback(50)
    temp_files.append(temp_sub_doc_path)

    sub_doc = Document(temp_sub_doc_path)

    max_chapter = get_max_main_chapter_number_from_toc(temp_main_doc_path)
    new_chapter_prefix = f"{max_chapter + 1}."

    modify_first_occurrence(sub_doc, "附件：复查结果", "附录：复查结果")
    if appendix_index is not None:
        modify_first_occurrence(sub_doc, "复查结果", f"复查结果{appendix_index}", paragraph_index=0)
    modify_first_occurrence(sub_doc, "关于本文档", "关于本附录", paragraph_index=1)
    update_section_references(sub_doc, r'(参见|参照)报告中(\d+\.\d+（1）节)', r'参见报告中' + new_chapter_prefix + r'\2')

    sub_doc.save(temp_sub_doc_path)
    progress_callback(70)

    word_app = None
    main_doc = None

    try:
        word_app = win32.Dispatch("Kwps.Application")
        word_app.Visible = False

        main_doc = word_app.Documents.Open(temp_main_doc_path)
        sub_doc = word_app.Documents.Open(temp_sub_doc_path)

        main_doc.Range(main_doc.Content.End - 1).InsertAfter('\n')

        sub_doc.Range().Copy()
        main_doc.Range(main_doc.Content.End - 1).Paste()

        first_paragraph = main_doc.Paragraphs(main_doc.Paragraphs.Count - sub_doc.Paragraphs.Count + 1)
        if first_paragraph.Range.ListFormat.ListString:
            first_paragraph.Range.ListFormat.RemoveNumbers()

        main_doc.SaveAs(temp_main_doc_path)
        logging.info(f"合并后的文档已保存到临时文件 {temp_main_doc_path}")
        progress_callback(90)

        if add_red_note and is_last:
            add_red_note_to_doc(temp_main_doc_path)
            progress_callback(100)

    except AttributeError as e:
        logging.error(f"发生属性错误: {e}")
    except Exception as e:
        logging.error(f"发生错误: {e}")
    finally:
        if sub_doc is not None and hasattr(sub_doc, 'Close'):
            logging.info("正在关闭 sub_doc")
            sub_doc.Close(False)
        if main_doc is not None and hasattr(main_doc, 'Close'):
            logging.info("正在关闭 main_doc")
            main_doc.Close(False)

def select_files(entry):
    file_paths = filedialog.askopenfilenames(filetypes=[("Word files", "*.docx")])
    entry.delete(0, tk.END)
    entry.insert(0, ';'.join(file_paths))

def select_folder(entry):
    folder_path = filedialog.askdirectory()
    if folder_path:
        entry.delete(0, tk.END)
        entry.insert(0, folder_path)

def check_admin_privileges(file_path):
    try:
        with open(file_path, 'a'):
            pass
        return True
    except PermissionError:
        return False

def main():
    root = tk.Tk()
    root.title("渗透测试报告文档合并工具 v6.1")

    tk.Label(root, text="主文档路径（初测）:").grid(row=0, column=0, padx=10, pady=10)
    main_doc_entry = tk.Entry(root, width=50)
    main_doc_entry.grid(row=0, column=1, padx=10, pady=10)
    tk.Button(root, text="浏览", command=lambda: select_files(main_doc_entry)).grid(row=0, column=2, padx=10, pady=10)

    tk.Label(root, text="副文档路径（复测）:").grid(row=1, column=0, padx=10, pady=10)
    sub_doc_entry = tk.Entry(root, width=50)
    sub_doc_entry.grid(row=1, column=1, padx=10, pady=10)
    tk.Button(root, text="浏览", command=lambda: select_files(sub_doc_entry)).grid(row=1, column=2, padx=10, pady=10)

    tk.Label(root, text="输出文档路径:").grid(row=2, column=0, padx=10, pady=10)
    output_doc_entry = tk.Entry(root, width=50)
    output_doc_entry.grid(row=2, column=1, padx=10, pady=10)
    tk.Button(root, text="浏览文件夹", command=lambda: select_folder(output_doc_entry)).grid(row=2, column=2, padx=5, pady=10)

    progress = ttk.Progressbar(root, orient="horizontal", length=400, mode="determinate")
    progress.grid(row=3, column=0, columnspan=4, padx=10, pady=20)

    def start_merge():
        main_doc_path = main_doc_entry.get()
        sub_doc_paths = sub_doc_entry.get().split(';')
        output_doc_path = output_doc_entry.get()

        if not main_doc_path:
            messagebox.showerror("错误", "请指定主文档路径！")
            return

        if os.path.isdir(output_doc_path):
            output_doc_path = os.path.join(output_doc_path, 'merged_output.docx')

        if not check_admin_privileges(output_doc_path):
            if messagebox.askyesno("权限错误", "输出文档路径需要管理员权限才能写入。是否尝试以管理员身份运行此操作？"):
                try:
                    ctypes.windll.shell32.ShellExecuteW(None, "runas", sys.executable, " ".join(sys.argv), None, 1)
                except Exception as e:
                    messagebox.showerror("错误", f"尝试获取管理员权限失败: {e}")
                return
            else:
                return

        def progress_callback(value):
            progress['value'] = value
            root.update_idletasks()

        temp_files = []
        temp_main_doc_path = main_doc_path.replace(".docx", "_temp.docx")

        remove_end_text(main_doc_path, temp_main_doc_path, "【全文结束】")
        progress_callback(10)
        temp_files.append(temp_main_doc_path)

        update_table_of_contents(temp_main_doc_path)
        progress_callback(30)

        for i, sub_doc_path in enumerate(sub_doc_paths):
            is_last = (i == len(sub_doc_paths) - 1)
            appendix_index = i + 1 if len(sub_doc_paths) > 1 else None
            merge_documents(temp_main_doc_path, sub_doc_path, temp_main_doc_path, progress_callback,
                            add_red_note=is_last, is_last=is_last, temp_files=temp_files, appendix_index=appendix_index)

        update_table_of_contents(temp_main_doc_path)
        progress_callback(95)

        if os.path.exists(output_doc_path):
            os.remove(output_doc_path)

        os.rename(temp_main_doc_path, output_doc_path)
        logging.info(f"最终文档已保存到 {output_doc_path}")
        progress_callback(100)

        for temp_file in temp_files:
            if os.path.exists(temp_file):
                os.remove(temp_file)
                logging.info(f"临时文件 {temp_file} 已删除")

        messagebox.showinfo("完成", "文档合并已完成")

    tk.Button(root, text="开始合并", command=start_merge).grid(row=4, column=1, columnspan=2, padx=10, pady=20)
    tk.Label(root, text="若有多次复测的，复测报告一定要严格按其复测时间的前后顺序依次合并！！！", fg='red', font=("宋体", 10)).grid(row=6, column=1, columnspan=2, padx=10, pady=10)

    root.mainloop()

if __name__ == "__main__":
    main()