from docx import Document
import logging

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def remove_end_text(doc_path, output_path, end_text="【全文结束】"):
    try:
        doc = Document(doc_path)
    except Exception as e:
        logging.error(f"打开文档 {doc_path} 时出错: {e}")
        raise

    if doc.paragraphs:
        last_paragraph = doc.paragraphs[-1]
        if end_text in last_paragraph.text:
            last_paragraph.text = last_paragraph.text.replace(end_text, "").strip()
            doc.save(output_path)
            logging.info(f"文档 {output_path} 已删除末尾的 {end_text}")
        else:
            logging.warning(f"文档 {doc_path} 的末尾不包含: {end_text}")
            doc.save(output_path)  # 仍然保存副本
    else:
        logging.warning(f"文档 {doc_path} 是空的")
        doc.save(output_path)
